"""Build a clean OCPP 1.6 Word guide with proper tables and no generator metadata."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

MD_PATH = Path(r"C:\Users\massi\Downloads\OCPP_Simulator\evse-ocpp16-console\docs\OCPP-1.6-Guide.md")
OUT_PATHS = [
    Path(r"C:\Users\massi\Downloads\OCPP_Simulator\evse-ocpp16-console\docs\OCPP-1.6-Guide.docx"),
    Path(r"C:\Users\massi\Downloads\OCPP_Simulator\evse-ocpp16-console\docs\OCPP-1.6-Guide-Detailed.docx"),
]

BRAND = RGBColor(0xC0, 0x24, 0x34)
HEADER_FILL = "C02434"
HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
ALT_ROW_FILL = "FBEAEA"


def strip_md(text: str) -> str:
    text = text.replace("**", "").replace("`", "")
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.strip()


def set_run_font(run, *, size=11, bold=False, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill_hex: str):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    for child in list(tc_pr):
        if child.tag == qn("w:shd"):
            tc_pr.remove(child)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold=False, color=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(strip_md(text))
    set_run_font(run, size=size, bold=bold, color=color)


def add_formatted_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    # normalize row widths
    norm = [list(r) + [""] * (cols - len(r)) for r in rows]

    table = doc.add_table(rows=len(norm), cols=cols)
    table.style = "Table Grid"
    table.autofit = True

    for i, row in enumerate(norm):
        for j, value in enumerate(row):
            cell = table.rows[i].cells[j]
            if i == 0:
                set_cell_shading(cell, HEADER_FILL)
                set_cell_text(cell, value, bold=True, color=HEADER_TEXT, size=9)
            else:
                if i % 2 == 0:
                    set_cell_shading(cell, ALT_ROW_FILL)
                set_cell_text(cell, value, bold=False, color=None, size=9)

    # Prefer full-width tables
    try:
        table.allow_autofit = True
    except Exception:
        pass

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def add_rich_paragraph(doc: Document, line: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    # split on **bold**
    parts = line.split("**")
    for idx, part in enumerate(parts):
        if not part:
            continue
        # also strip inline backticks visually
        chunks = part.split("`")
        for c_idx, chunk in enumerate(chunks):
            if not chunk:
                continue
            run = p.add_run(chunk)
            set_run_font(
                run,
                size=11,
                bold=(idx % 2 == 1),
                name="Consolas" if c_idx % 2 == 1 else "Calibri",
            )


def clear_metadata(doc: Document):
    cp = doc.core_properties
    cp.author = "Massive Mobility"
    cp.last_modified_by = "Massive Mobility"
    cp.comments = ""
    cp.subject = "OCPP 1.6 practical guide"
    cp.keywords = "OCPP, OCPP 1.6, Charge Point, CSMS, Massive Mobility"
    cp.category = "Technical documentation"
    cp.title = "OCPP 1.6 — Practical Guide"
    # Remove revision/generator leftovers if present in app props
    try:
        app = doc.part.package.part_related_by(
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties"
        )
    except Exception:
        app = None


def build_document(md_text: str) -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    in_code = False
    code_lines: list[str] = []
    table_rows: list[list[str]] = []

    def flush_code():
        nonlocal code_lines
        if not code_lines:
            return
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.3)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run("\n".join(code_lines))
        set_run_font(run, size=9, name="Consolas", color=RGBColor(0x33, 0x33, 0x33))
        code_lines = []

    def flush_table():
        nonlocal table_rows
        if not table_rows:
            return
        add_formatted_table(doc, table_rows)
        table_rows = []

    for raw in md_text.splitlines():
        line = raw.rstrip()

        if line.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_table()
                in_code = True
                code_lines = []
            continue

        if in_code:
            code_lines.append(line)
            continue

        if line.startswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            if cells and all(set(c) <= set("-: ") for c in cells):
                continue
            table_rows.append(cells)
            continue

        flush_table()

        if not line or line.strip() == "---":
            continue

        if line.startswith("# "):
            h = doc.add_heading(strip_md(line[2:]), level=0)
            for run in h.runs:
                set_run_font(run, size=22, bold=True, color=BRAND, name="Calibri")
        elif line.startswith("## "):
            h = doc.add_heading(strip_md(line[3:]), level=1)
            for run in h.runs:
                set_run_font(run, size=14, bold=True, color=RGBColor(0x11, 0x18, 0x27), name="Calibri")
        elif line.startswith("### "):
            h = doc.add_heading(strip_md(line[4:]), level=2)
            for run in h.runs:
                set_run_font(run, size=12, bold=True, color=RGBColor(0x37, 0x41, 0x51), name="Calibri")
        elif line.startswith("- "):
            p = doc.add_paragraph(strip_md(line[2:]), style="List Bullet")
            for run in p.runs:
                set_run_font(run, size=11)
        elif line.startswith("*") and line.endswith("*") and not line.startswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(strip_md(line.strip("* ")))
            set_run_font(run, size=10)
            run.italic = True
        else:
            add_rich_paragraph(doc, line)

    flush_table()
    flush_code()
    clear_metadata(doc)
    return doc


def main():
    md = MD_PATH.read_text(encoding="utf-8")
    # Ensure footer has no generator / AI wording
    md = re.sub(r"(?i)\b(ai[- ]generated|generated by .+|chatgpt|openai|claude|llm)\b", "", md)

    doc = build_document(md)
    wrote = []
    for path in OUT_PATHS:
        try:
            doc.save(path)
            wrote.append(str(path))
        except PermissionError:
            fallback = path.with_name(path.stem + "-clean.docx")
            doc.save(fallback)
            wrote.append(str(fallback) + " (fallback)")
    print("wrote:")
    for w in wrote:
        print(" ", w)


if __name__ == "__main__":
    main()
