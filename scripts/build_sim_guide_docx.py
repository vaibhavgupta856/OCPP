from docx import Document
from pathlib import Path

md_path = Path(r"C:\Users\massi\Downloads\OCPP_Simulator\evse-ocpp16-console\docs\Massive-Mobility-Simulator-Guide.md")
out_path = Path(r"C:\Users\massi\Downloads\OCPP_Simulator\evse-ocpp16-console\docs\Massive-Mobility-Simulator-Guide.docx")
legacy_names = [
    "Pier-Simulator-Guide.docx",
    "~$er-Simulator-Guide.docx",
]

text = md_path.read_text(encoding="utf-8")
doc = Document()
in_code = False
for raw in text.splitlines():
    line = raw.rstrip()
    if line.startswith("```"):
        in_code = not in_code
        continue
    if not line:
        continue
    if in_code:
        doc.add_paragraph(line)
        continue
    if line.startswith("# "):
        doc.add_heading(line[2:].strip(), level=0)
    elif line.startswith("## "):
        doc.add_heading(line[3:].strip(), level=1)
    elif line.startswith("### "):
        doc.add_heading(line[4:].strip(), level=2)
    elif line.startswith("|"):
        cells = [c.strip() for c in line.strip("|").split("|")]
        if all(set(c) <= set("-: ") for c in cells):
            continue
        doc.add_paragraph(" | ".join(cells))
    elif line.startswith("- "):
        doc.add_paragraph(line[2:].strip(), style="List Bullet")
    elif line.startswith("*") and line.endswith("*"):
        doc.add_paragraph(line.strip("* "))
    else:
        doc.add_paragraph(line)

doc.save(out_path)
print("wrote", out_path)

docs = out_path.parent
for name in legacy_names:
    p = docs / name
    if p.exists():
        try:
            p.unlink()
            print("removed legacy", name)
        except Exception as e:
            print("could not remove", name, "(close Word if open):", e)
